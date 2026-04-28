"""
Genera un documento Word con el análisis completo de la Crisis de los Tulipanes
"""

from docx import Document
from docx.shared import Inches, Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

# Crear documento
doc = Document()

# ============================================================
# TÍTULO PRINCIPAL
# ============================================================
title = doc.add_heading('ANÁLISIS DE LA CRISIS DE LOS TULIPANES', 0)
title.alignment = WD_ALIGN_PARAGRAPH.CENTER

# Subtítulo
subtitle = doc.add_paragraph('Estudio Histórico y Análisis de Datos (1634-1638)')
subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER

doc.add_paragraph()

# ============================================================
# RESUMEN EJECUTIVO
# ============================================================
doc.add_heading('RESUMEN EJECUTIVO', level=1)

resumen = doc.add_paragraph(
    '''La Crisis de los Tulipanes (1635-1637) constituye la primera burbuja especulativa documentada en la historia económica. '''
    '''Este análisis examina el comportamiento del mercado holandés de tulipanes durante este período, '''
    '''utilizando un dataset histórico que captura la evolución de precios, volúmenes y actores del mercado.'''
)

# Estadísticas clave en tabla
doc.add_heading('Estadísticas Clave', level=2)

table = doc.add_table(rows=5, cols=2)
table.style = 'Table Grid'

stats_data = [
    ('Período analizado', 'Enero 1634 - Junio 1638'),
    ('Precio máximo alcanzado', '4,000 guilders (Febrero 1637)'),
    ('Aumento máximo', '800x (de 5 a 4,000 guilders)'),
    ('Caída máxima', '-95% (en cuestión de días)'),
    ('Volumen total negociado', '36,355 bulbos'),
]

for i, (metric, value) in enumerate(stats_data):
    row = table.rows[i]
    row.cells[0].text = metric
    row.cells[1].text = value

doc.add_paragraph()

# ============================================================
# CAPÍTULO 1: LIMPIEZA DE DATOS
# ============================================================
doc.add_heading('1. LIMPIEZA DE DATOS', level=1)

doc.add_paragraph(
    '''Se realizó una limpieza del dataset original para corregir problemas de formato '''
    '''y asegurar la consistencia de los datos.'''
)

doc.add_heading('Problemas Detectados y Correcciones', level=2)

# Tabla de problemas
table2 = doc.add_table(rows=4, cols=3)
table2.style = 'Table Grid'

# Headers
headers = ['Problema', 'Ubicación', 'Corrección']
for i, header in enumerate(headers):
    table2.rows[0].cells[i].text = header

problems = [
    ('Espacios en valores', 'valor intrínseco', 'Agregar comillas'),
    ('Valores con tildes', 'nueva_demanda', 'Corregir encoding'),
    ('Formato CSV', 'Líneas 35, 47', 'Verificar comas'),
]

for i, (prob, loc, fix) in enumerate(problems, 1):
    table2.rows[i].cells[0].text = prob
    table2.rows[i].cells[1].text = loc
    table2.rows[i].cells[2].text = fix

doc.add_paragraph()

# ============================================================
# CAPÍTULO 2: ANÁLISIS DE DATOS
# ============================================================
doc.add_heading('2. ANÁLISIS DE DATOS', level=1)

doc.add_heading('2.1 Estadísticas Generales', level=2)

table3 = doc.add_table(rows=6, cols=2)
table3.style = 'Table Grid'

general_stats = [
    ('Dataset final', '46 registros'),
    ('Columnas', '11'),
    ('Período', '4.5 años'),
    ('Precio mínimo', '5 guilders'),
    ('Precio máximo', '4,000 guilders'),
    ('Variedades analizadas', '3 (Semper Augustus, General, Admiral)'),
]

for i, (metric, value) in enumerate(general_stats):
    table3.rows[i].cells[0].text = metric
    table3.rows[i].cells[1].text = value

doc.add_paragraph()

# ============================================================
# Análisis por Variedad
# ============================================================
doc.add_heading('2.2 Análisis por Variedad', level=2)

doc.add_paragraph(
    '''Se analizaron las tres variedades principales de tulipanes presentes en el dataset:'''
)

# Tabla de variedades
table4 = doc.add_table(rows=4, cols=3)
table4.style = 'Table Grid'

variedad_headers = ['Variedad', 'Volumen Total', 'Precio Promedio']
for i, header in enumerate(variedad_headers):
    table4.rows[0].cells[i].text = header

variedades = [
    ('Semper Augustus', '14,050 bulbos (56%)', '~800 guilders'),
    ('General', '8,150 bulbos (33%)', '~600 guilders'),
    ('Admiral', '2,900 bulbos (11%)', '~400 guilders'),
]

for i, (var, vol, precio) in enumerate(variedades, 1):
    table4.rows[i].cells[0].text = var
    table4.rows[i].cells[1].text = vol
    table4.rows[i].cells[2].text = precio

doc.add_paragraph()

# Análisis TOP y menos vendidos
doc.add_heading('TOP 5: Tulipanes Más Vendidos', level=3)
doc.add_paragraph('1. Semper Augustus: 56% del volumen total')
doc.add_paragraph('2. General: 33% del volumen total')
doc.add_paragraph('3. Admiral: 11% del volumen total')

doc.add_heading('TOP 5: Tulipanes Menos Vendidos', level=3)
doc.add_paragraph('1. Admiral: 11% del volumen total (menos exclusivo)')
doc.add_paragraph('2. General: 33% (no era la variedad "rara")')

doc.add_paragraph()
doc.add_paragraph(
    'NOTA: No hubo variedades realmente impopulares. La diferencia es que Admiral tuvo menor volumen '
    'por ser menos exclusiva que Semper Augustus.'
)

# ============================================================
# Análisis por Fase Económica
# ============================================================
doc.add_heading('2.3 Análisis por Fase Económica', level=2)

table5 = doc.add_table(rows=8, cols=4)
table5.style = 'Table Grid'

phase_headers = ['Fase', 'Período', 'Precio Inicio-Fin', 'Cambio']
for i, header in enumerate(phase_headers):
    table5.rows[0].cells[i].text = header

phases = [
    ('Incipiente', '1634', '5 → 12', '+140%'),
    ('Boom', '1635', '12 → 250', '+1,983%'),
    ('Pico', '1636', '400 → 3,200', '+700%'),
    ('Pico Máximo', 'Feb 1637', '3,500 → 4,000', '+14%'),
    ('Crash', 'Feb 1637', '4,000 → 200', '-95%'),
    ('Recovery', '1637-1638', '200 → 8', '-96%'),
    ('Estabilizado', 'Jun 1638', '8 → 8', '0%'),
]

for i, (fase, periodo, cambio, pct) in enumerate(phases, 1):
    table5.rows[i].cells[0].text = fase
    table5.rows[i].cells[1].text = periodo
    table5.rows[i].cells[2].text = cambio
    table5.rows[i].cells[3].text = pct

doc.add_paragraph()

# ============================================================
# Análisis por Tipo de Comprador
# ============================================================
doc.add_heading('2.4 Análisis por Tipo de Comprador', level=2)

table6 = doc.add_table(rows=6, cols=3)
table6.style = 'Table Grid'

comprador_headers = ['Tipo', 'Fase Activa', 'Comportamiento']
for i, header in enumerate(comprador_headers):
    table6.rows[0].cells[i].text = header

compradores = [
    ('Coleccionista', '1634-1635, 1637-1638', 'Compra por valor intrínseco'),
    ('Merchant', 'Todo el período', 'Distribuidor, busca liquidez'),
    ('Especulador', '1635-1637', 'Apuesta a precio futuro, alto riesgo'),
    ('Granjero', '1636', 'Entró tarde, fue el más dañado'),
    ('Comerchant', '1637', 'Compra en pánico, busca ganancias'),
]

for i, (tipo, fase, comp) in enumerate(compradores, 1):
    table6.rows[i].cells[0].text = tipo
    table6.rows[i].cells[1].text = fase
    table6.rows[i].cells[2].text = comp

doc.add_paragraph()

# ============================================================
# CAPÍTULO 3: EXPLICACIÓN DETALLADA
# ============================================================
doc.add_heading('3. EXPLICACIÓN DETALLADA DE LA CRISIS', level=1)

# Fase 1
doc.add_heading('Fase 1: Incipiente (1634)', level=2)
doc.add_paragraph(
    '''Los tulipanes Semper Augustus llegan de Turquía con patrones únicos (colores striados). '''
    '''Los coleccionistas de la corte holandesa comienzan a pagar precios elevados. '''
    '''No hay especulación — es un artículo de lujo.'''
)
doc.add_paragraph('Precio: 5 → 12 guilders (+140%)')
doc.add_paragraph('Volumen: 20 → 80 bulbos')
doc.add_paragraph('Perfil: coleccionistas enriquecidos')

# Fase 2
doc.add_heading('Fase 2: El Boom (1635)', level=2)
doc.add_paragraph(
    '''La noticia se extiende. Los Forward Contracts nacen — se puede comprar tulipanes '''
    '''que aún no existen, solo con un contrato. Los especuladores entran.'''
)
doc.add_paragraph('Precio: 12 → 250 guilders (+1,983%)')
doc.add_paragraph('Volumen: 80 → 1,000 bulbos')
doc.add_paragraph('Perfil: coleccionistas + primeros especuladores')

# Fase 3
doc.add_heading('Fase 3: El Pico (1636)', level=2)
doc.add_paragraph(
    '''Manía colectiva. Todos hablan de tulipanes. Granjeros, sirvientes, comerciantes '''
    '''venden sus pertenencias para comprar bulbos. Los precios suben a 3,200 guilders.'''
)
doc.add_paragraph('Precio: 400 → 3,200 guilders (+700%)')
doc.add_paragraph('Volumen: 1,200 → 1,000 bulbos (menos oferta = más precio)')
doc.add_paragraph('Perfil: especuladores dominan (70%)')

# Fase 4
doc.add_heading('Fase 4: El Crash (Febrero 1637)', level=2)
doc.add_paragraph(
    '''Pánico. Alguien intenta vender y no encuentra comprador. '''
    '''El rumor se extiende. En 3 días, el precio cae 95%.'''
)
doc.add_paragraph('Feb 1637 inicio: 4,000 guilders')
doc.add_paragraph('Feb 1637 medio: 500 guilders')
doc.add_paragraph('Feb 1637 fin: 200 guilders')
doc.add_paragraph('Caída: -95% en días')
doc.add_paragraph('Los granjeros pierden todo — habían invertido sus ahorros.')

# Fase 5
doc.add_heading('Fase 5: Recovery (1637-1638)', level=2)
doc.add_paragraph(
    '''El mercado bottoms en 5-8 guilders (valor real del bulbo). '''
    '''Los coleccionistas vuelven a comprar a precios razonables. '''
    '''El mercado se estabiliza.'''
)
doc.add_paragraph('Precio: 200 → 8 guilders (-96%)')
doc.add_paragraph('Volumen: 50 → 800 bulbos (recuperación)')
doc.add_paragraph('Perfil: coleccionistas de vuelta')

doc.add_paragraph()

# ============================================================
# CAPÍTULO 4: CAUSAS DE LA CRISIS
# ============================================================
doc.add_heading('4. CONCLUSIÓN: ¿POR QUÉ OCURRIÓ LA CRISIS?', level=1)

causas = [
    ('1. Activos ilíquidos', 'Los tulipanes no podían revenderse fácilmente — el mercado era opaco'),
    ('2. Contratos forward', 'Compras sin entrega inmediata, pura especulación'),
    ('3. Efecto manada', 'Todo el mundo quería participar'),
    ('4. Información asimétrica', 'Nadie sabía el verdadero valor'),
    ('5. Apalancamiento', 'Los contratos permitían apuesta enorme sin capital'),
    ('6. Falta de regulación', 'No existían protecciones al inversionista'),
]

for causa, desc in causas:
    p = doc.add_paragraph()
    p.add_run(causa).bold = True
    p.add_run(': ' + desc)

doc.add_paragraph()

# ============================================================
# INSIGHTS CLAVE
# ============================================================
doc.add_heading('5. INSIGHTS CLAVE DEL ANÁLISIS', level=1)

insights = [
    'Semper Augustus dominó el mercado — variedad premium con más volumen y precio',
    'El crash fue VERTICAL — de 4,000 a 100 en un mes (no gradual)',
    'Especuladores dominante — 70% de transacciones en fase pico',
    'Granjeros perdieron todo — entraron tarde, vendieron en pánico',
    'Recovery tomó 2 años — de 1637 a 1638 para estabilizarse',
]

for i, insight in enumerate(insights, 1):
    doc.add_paragraph(f'{i}. {insight}')

doc.add_paragraph()

# ============================================================
# VISUALIZACIONES GENERADAS
# ============================================================
doc.add_heading('6. VISUALIZACIONES GENERADAS', level=1)

doc.add_paragraph(
    '''Se generaron 4 visualizaciones en formato PNG que acompañan este análisis:'''
)

viz = [
    'grafico1_precio_vs_tiempo.png — Evolución del precio en el tiempo con colores por fase',
    'grafico2_variedad.png — Pie y Bar de tulipanes por variedad',
    'grafico3_fases.png — 4 subgráficos de análisis por fases',
    'grafico4_calor.png — Heatmap precio x mes x año',
]

for v in viz:
    doc.add_paragraph(v)

doc.add_paragraph()

# ============================================================
# NOTA FINAL
# ============================================================
doc.add_heading('NOTA', level=1)
doc.add_paragraph(
    '''Este análisis es una aproximación histórica basada en datos simulados para fines educativos. '''
    '''Los valores exactos de la Crisis de los Tulipanes han sido objeto de debate entre historiadores, '''
    '''pero los patrones generales identificados reflejan el comportamiento documentado de la burbuja.'''
)

# Guardar documento
doc.save('analisis_crisis_tulipanes.docx')
print("Documento creado: analisis_crisis_tulipanes.docx")